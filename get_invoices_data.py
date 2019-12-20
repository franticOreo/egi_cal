import docx
import re, os
import csv
from datetime import datetime
import pickle

# doc = docx.Document('Copy of Tax Invoice for Egons GC Dec.docx')

def main():
    # get_cust_info()
    list_of_cust = get_cust_info()
    clean_cust_info(list_of_cust)


def get_cust_info():
    '''
    This function iterates of the Invoices directory, finds the document files.
    It then creates a docx Document object of the file in order to read the text.
    It processes this object by splitting the first paragraph to source majority of
    the customer information. This information is then appended to a list of cutomers
    list.
    '''

    inv_dir = r'/home/franticoreo/egi_cal/Invoices/'
    list_of_cust = []




    for filename in os.listdir(inv_dir):
        if filename.endswith('.docx'):
            doc = docx.Document(inv_dir + filename)
            fp = doc.paragraphs[0].text



            # get date of work from first paragraph
            # date = get_date(fp, filename) 
            rate = get_rate_from_table(doc, filename)


            try:
                # get a string of customer info after "To :" in invoice
                after_to = fp.split('To :')[1]
                # remove words after the word "Scope"
                customer_info = after_to.split('Scope')[0]
                customer_info = customer_info.split(',')
                customer_info.append(rate)
                # customer_info.append(date)
                list_of_cust.append(customer_info)

                # print(list_of_cust)
                # list_of_cust.append(rate)

            except:
                print('Error handling Invoice: '+ filename)
    # print(list_of_cust)
    return list_of_cust


def get_cust_email(name):
    with open('/home/franticoreo/egi_cal/docs/clean_entries.txt', 'rb') as file:
        subj_recip = pickle.load(file)
        # print(subj_recip)
    # read through list of dicts
    # if the dict recip starts with same as name return email
    for email_d in subj_recip:
        if email_d['recip'].startswith(name):
            # cleaning to extract email from <> tags
            email = email_d['recip'].split('<')[1].split('>')[0]
            return email


def get_date(filename):
    """Creates a datetime object from date text in invoice paragraph

    Args:
        filename: string
    Returns:
        datetime object
    """
    if filename.endswith('.docx'):
        doc = docx.Document(inv_dir + filename)
        paragraph = doc.paragraphs[0].text
    try:
        date = paragraph.split('From:')[0].split('Date:')[1].strip()
        if len(date) > 8:
            date = datetime.strptime(date, '%d/%m/%Y') 
            # print(date)
            return date

        else:
            date = datetime.strptime(date, '%d/%m/%y') 
            # print(date)
            return date
    
    except Exception as e:
        print('Error finding date for: ', filename)
        print(e)

def get_last_time_worked(filename, customer_list):
    """Gets the last date of work for each customer in list
    
    Args:
        filename: string
        customer_list: list
    Returns:
        last_work_date: dictionary {name, datetime}
    """
    customer_list = clean_cust_info(customer_list)





    

def get_rate_from_table(doc, filename):
    '''
    This function gets the document object, cleans the data to convert to an integer.
    The data is cleaned and sorted, setting empty data to type None. If the data is of
    the type float, the rate is computed. A lambda function is used to standardise the 
    rate to be either 35 or 50.
    '''
    print(filename)
    table = doc.tables[0]

    # get quantity of hours
    quant_hours = table.columns[0].cells[1].text
    cost_hours = table.columns[2].cells[1].text

    # cleaning tax invoice data
    quant_hours = quant_hours.replace('hrs', '')
    quant_hours = quant_hours.replace('hr', '')

    cost_hours = cost_hours.replace('$', '')
    cost_hours = cost_hours.replace('.', '')

    if quant_hours == '' or cost_hours == '':
        print('Empty Values from Tax Invoice: ' + filename)
        return None

    try:
        hourly_rate = float(cost_hours) / float(quant_hours)
        # print(hourly_rate)
        standard_rates = [35, 50]
        # return whether the rate is closer to $35 or $50 to standardise rates
        return min(standard_rates, key=lambda x: abs(x - hourly_rate))

    except Exception as e:
        print('Error Computing rate for: '+ filename)
        print("quant hours" + quant_hours)
        print("cost hours" + cost_hours)
        # print(e)
        return None


    print(total_cost, hours)
    return (total_cost, hours)

def clean_cust_info(customer_list):
    '''
    This function recieves a list of customer arrays , if the array fits
    the standard structure [name, address, suburb] it is then cleaned by removing
    white space and concatenating the address and suburb. These arrays are then converted
    into tuples to be able to added to a set. Therefore, the set removes repeated entries.
    '''

    unique_cust = set()
    dirty_cust = []

    for customer_array in customer_list:
        try:
            name = customer_array[0].strip()
            email = get_cust_email(name)
            # remove trailing whitespace from suburb item with rstrip method
            suburb = str(customer_array[2]).rstrip()
            address = str(customer_array[1]) 
            rate = str(customer_array[-1])
            # print(customer_array)
            # get customer entry, create a tuple(unhashable) to the set
            l = (name, address, suburb, rate, email)
            unique_cust.add(l)

        except Exception as e:
            print('Error handling Customer: '+ str(customer_array))
            print(e)
            dirty_cust.append(customer_array)
            
    print('Dirty Customers:')
    print(dirty_cust)

    print('Clean, Unique Customers:')
    print()
    print(unique_cust)
    return unique_cust




if __name__ == '__main__':
    main()
