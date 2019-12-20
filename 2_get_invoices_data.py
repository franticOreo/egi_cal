import docx
import re, os
import csv
from datetime import datetime
import pickle
from pprint import pprint



def main():
    # get_cust_info()
    list_of_cust = get_data_from_invoices()
    # unique_customers(list_of_cust)
    # clean_cust_info(list_of_cust)


def get_data_from_invoices():
    '''
    This function iterates of the Invoices directory, finds the document files.
    It then creates a docx Document object of the file in order to read the text.
    It processes this object by splitting the first paragraph to source majority of
    the customer information. This information is then appended to a list of cutomers
    list.
    '''

    inv_dir = r'/home/franticoreo/egi_cal/Invoices/'
    list_of_inv = []
    inv_with_error = []

    for filename in os.listdir(inv_dir):
        if filename.endswith('.docx'):
            doc = docx.Document(inv_dir + filename)
            fp = doc.paragraphs[0].text

            # get date of work from first paragraph
            date = get_date(fp, filename)
            rate = get_rate_from_table(doc, filename)

            invoice_d = {}
            try:
                # get a string of customer info after "To :" in invoice
                after_to = fp.split('To :')[1]
                # remove words after the word "Scope"
                customer_info = after_to.split('Scope')[0]
                # creates an array [name, address, suburb] by splitting by a comma
                customer_info = customer_info.split(',')
                invoice_d['name'] = customer_info[0].strip()
                invoice_d['email'] = get_cust_email(invoice_d['name'])
                invoice_d['address'] = customer_info[1]
                invoice_d['suburb'] = customer_info[2].rstrip()
                invoice_d['rate'] = rate
                invoice_d['date'] = date
                # customer_info.append(date)
                list_of_inv.append(invoice_d)

            except:
                print('Error handling Invoice: '+ filename)
                inv_with_error.append(filename)
    print(list_of_inv)
    print('Following invoices were unable to extracted...')
    pprint(inv_with_error)

    return list_of_inv




def get_cust_email(name):
    with open('docs/clean_entries.txt', 'rb') as file:
        subj_recip = pickle.load(file)
        # print(subj_recip)
    # read through list of dicts
    # if the dict recip starts with same as name return email
    for email_d in subj_recip:
        if email_d['recip'].startswith(name):
            # cleaning to extract email from <> tags
            email = email_d['recip'].split('<')[1].split('>')[0]
            return email


def get_date(paragraph, filename):
    """Creates a datetime object from date text in invoice paragraph

    Args:
        filename: string
    Returns:
        datetime object
    """
    try:
        date = paragraph.split('From:')[0].split('Date:')[1].strip()
        if len(date) > 8:
            date = datetime.strptime(date, '%d/%m/%Y') 
            # print(date)
            return date

        else:
            date = datetime.strptime(date, '%d/%m/%y') 
            # print(date)
            return date
    
    except Exception as e:
        print('Error finding date for: ', filename)
        print(e)

def get_last_time_worked():
    """Gets the last date of work for each customer in list

    """
    inv_list = get_data_from_invoices()
    # go through invoice data get customer name and date
    # if customer seen, check if date is later
    cust_and_date = []

    for inv_d in inv_list:
        cust_and_date['name'] = inv_d['name']
        cust_and_date['date'] = inv_d['date']
        

    

def get_rate_from_table(doc, filename):
    '''
    This function gets the document object, cleans the data to convert to an integer.
    The data is cleaned and sorted, setting empty data to type None. If the data is of
    the type float, the rate is computed. A lambda function is used to standardise the 
    rate to be either 35 or 50.
    '''
    table = doc.tables[0]

    # get quantity of hours
    quant_hours = table.columns[0].cells[1].text
    cost_hours = table.columns[2].cells[1].text

    # cleaning tax invoice data
    quant_hours = quant_hours.replace('hrs', '')
    quant_hours = quant_hours.replace('hr', '')

    cost_hours = cost_hours.replace('$', '')
    cost_hours = cost_hours.replace('.', '')

    if quant_hours == '' or cost_hours == '':
        print('Empty Values from Tax Invoice: ' + filename)
        return None

    try:
        hourly_rate = float(cost_hours) / float(quant_hours)
        # print(hourly_rate)
        standard_rates = [35, 50]
        # return whether the rate is closer to $35 or $50 to standardise rates
        return min(standard_rates, key=lambda x: abs(x - hourly_rate))

    except Exception as e:
        print('Error Computing rate for: '+ filename)
        print("quant hours" + quant_hours)
        print("cost hours" + cost_hours)
        # print(e)
        return None


    print(total_cost, hours)
    return (total_cost, hours)


def unique_customers(list_of_inv):
    unique_cust = []


    for cust_d in list_of_inv:
        # remove date to only retain customer info
        del cust_d['date']
        if cust_d in unique_cust:
            continue
        else:
            unique_cust.append(cust_d)

    print('Clean, Unique Customers:')
    print()
    print(unique_cust)
    print(len(unique_cust))
    return unique_cust

if __name__ == '__main__':
    main()
